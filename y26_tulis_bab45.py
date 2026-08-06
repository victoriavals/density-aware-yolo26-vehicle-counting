"""Pembangkit naskah BAB IV & BAB V (Microsoft Word) — Tesis Naufal Firdaus.

Menghasilkan `TESIS_BAB4-5.docx` yang formatnya diselaraskan byte-per-byte dengan naskah
otoritatif `TESIS_BAB1-3_REVISI_SIDANG_v7.docx`: A4, margin 3-3-4-3 cm, Times New Roman
12 pt, spasi baris 1,15, indentasi baris pertama 1 cm, keterangan tabel DI ATAS tabel dan
keterangan gambar DI BAWAH gambar, gaya `_tabel`/`_gambar` 11 pt rata tengah.

Prinsip: SELURUH angka dibaca dari berkas hasil (`hasil_bab4_5/`, `eval_out/`, `logs/`).
Tidak ada angka yang diketik ulang di dalam prosa — bila sumber berubah, jalankan ulang.

    ./.venv/Scripts/python.exe y26_tulis_bab45.py

Penomoran melanjutkan naskah BAB 1-3: Tabel 4.1 dst (BAB 3 berhenti di Tabel 3.10),
Gambar 4.1 dst (BAB 3 berhenti di Gambar 3.7), sitasi [31] dst (BAB 1-3 memakai [1]-[30]).
"""
from __future__ import annotations

import csv
import json
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parent
HB = ROOT / "hasil_bab4_5"
OUT = ROOT / "TESIS_BAB4-5.docx"

FONT = "Times New Roman"
SPASI = 1.15
INDENT = Cm(1.0)

# ---------------------------------------------------------------- pembaca data
def baca_csv(rel: str) -> list[dict]:
    with open(HB / rel, encoding="utf-8-sig", newline="") as fh:
        return list(csv.DictReader(fh))


def baca_json(rel: str):
    return json.loads((HB / rel).read_text(encoding="utf-8"))


def ind(x, desimal: int = 2, persen: bool = False) -> str:
    """Format angka gaya Indonesia (koma desimal). persen=True mengalikan 100."""
    if x is None or x == "":
        return "-"
    v = float(x) * (100 if persen else 1)
    return f"{v:.{desimal}f}".replace(".", ",")


# --------------------------------------------------------------- gaya dokumen
def siapkan_dokumen() -> Document:
    doc = Document()
    s = doc.sections[0]
    s.page_width, s.page_height = Cm(21), Cm(29.7)
    s.top_margin = s.bottom_margin = s.right_margin = Cm(3)
    s.left_margin = Cm(4)

    n = doc.styles["Normal"]
    n.font.name = FONT
    n.font.size = Pt(12)
    n.font.color.rgb = RGBColor(0, 0, 0)
    n.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    pf = n.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.line_spacing = SPASI
    pf.first_line_indent = INDENT
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)

    for nama, ukuran, before, after, align in (
            ("Heading 1", 14, 18, 4, WD_ALIGN_PARAGRAPH.CENTER),
            ("Heading 2", 12, 8, 4, WD_ALIGN_PARAGRAPH.LEFT),
            ("Heading 3", 12, 8, 4, WD_ALIGN_PARAGRAPH.LEFT)):
        st = doc.styles[nama]
        st.font.name = FONT
        st.font.size = Pt(ukuran)
        st.font.bold = True
        st.font.italic = False
        st.font.color.rgb = RGBColor(0, 0, 0)
        st.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        p = st.paragraph_format
        p.alignment = align
        p.space_before, p.space_after = Pt(before), Pt(after)
        p.line_spacing = SPASI
        p.first_line_indent = Cm(0)
        p.keep_with_next = True

    for nama, before, after in (("_tabel", 8, 6), ("_gambar", 6, 6), ("_rumus", 6, 6)):
        st = doc.styles.add_style(nama, 1)          # 1 = WD_STYLE_TYPE.PARAGRAPH
        st.base_style = doc.styles["Normal"]
        st.font.name = FONT
        st.font.size = Pt(11)
        st.font.bold = False
        st.font.color.rgb = RGBColor(0, 0, 0)
        p = st.paragraph_format
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.line_spacing = SPASI
        p.first_line_indent = Cm(0)
        p.space_before, p.space_after = Pt(before), Pt(after)
    doc.styles["_tabel"].paragraph_format.keep_with_next = True
    return doc


# ------------------------------------------------------------- penulis blok
_MIRING = re.compile(r"\*(.+?)\*")


def tulis_runs(p, teks: str, size: Pt | None = None, bold: bool = False):
    """Tulis teks; potongan di antara *bintang* dicetak miring (istilah asing)."""
    pos = 0
    for m in _MIRING.finditer(teks):
        if m.start() > pos:
            r = p.add_run(teks[pos:m.start()])
            r.bold = bold
            if size:
                r.font.size = size
        r = p.add_run(m.group(1))
        r.italic = True
        r.bold = bold
        if size:
            r.font.size = size
        pos = m.end()
    if pos < len(teks):
        r = p.add_run(teks[pos:])
        r.bold = bold
        if size:
            r.font.size = size


def par(doc, teks: str):
    p = doc.add_paragraph()
    tulis_runs(p, teks)
    return p


def h1(doc, baris1: str, baris2: str):
    p = doc.add_paragraph(style="Heading 1")
    p.add_run(baris1)
    p.add_run().add_break(WD_BREAK.LINE)
    p.add_run(baris2)
    return p


def h2(doc, teks: str):
    p = doc.add_paragraph(style="Heading 2")
    tulis_runs(p, teks)
    return p


def h3(doc, teks: str):
    p = doc.add_paragraph(style="Heading 3")
    tulis_runs(p, teks)
    return p


def _border(el, sisi: str):
    b = OxmlElement(f"w:{sisi}")
    b.set(qn("w:val"), "single")
    b.set(qn("w:sz"), "4")
    b.set(qn("w:color"), "000000")
    el.append(b)


def tabel(doc, header: list[str], baris: list[list], lebar: list[float] | None = None,
          size_head: int = 10, size_body: int = 9, align_kanan: set[int] | None = None):
    """Tabel bergaris; kepala tebal rata tengah, isi 9 pt. lebar dalam cm."""
    t = doc.add_table(rows=1, cols=len(header))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    tblPr = t._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for sisi in ("top", "left", "bottom", "right", "insideH", "insideV"):
        _border(borders, sisi)
    tblPr.append(borders)

    align_kanan = align_kanan or set()
    for i, teks in enumerate(header):
        c = t.rows[0].cells[i]
        c.paragraphs[0].style = doc.styles["Normal"]
        pf = c.paragraphs[0].paragraph_format
        pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf.first_line_indent = Cm(0)
        pf.line_spacing = SPASI
        pf.space_after = Pt(0)
        tulis_runs(c.paragraphs[0], teks, size=Pt(size_head), bold=True)
    for b in baris:
        row = t.add_row()
        for i, teks in enumerate(b):
            c = row.cells[i]
            pf = c.paragraphs[0].paragraph_format
            pf.alignment = (WD_ALIGN_PARAGRAPH.RIGHT if i in align_kanan
                            else WD_ALIGN_PARAGRAPH.LEFT if i == 0
                            else WD_ALIGN_PARAGRAPH.CENTER)
            pf.first_line_indent = Cm(0)
            pf.line_spacing = SPASI
            pf.space_after = Pt(0)
            tulis_runs(c.paragraphs[0], str(teks), size=Pt(size_body))
    if lebar:
        for r in t.rows:
            for i, w in enumerate(lebar):
                r.cells[i].width = Cm(w)
    return t


class Nomor:
    """Penomoran otomatis tabel/gambar per bab, melanjutkan naskah BAB 1-3."""

    def __init__(self, bab: int):
        self.bab, self.t, self.g = bab, 0, 0
        self.daftar_t: list[tuple[str, str]] = []
        self.daftar_g: list[tuple[str, str]] = []

    def tabel(self, doc, judul: str) -> str:
        self.t += 1
        no = f"{self.bab}.{self.t}"
        p = doc.add_paragraph(style="_tabel")
        tulis_runs(p, f"Tabel {no} {judul}", size=Pt(11))
        self.daftar_t.append((no, judul))
        return no

    def gambar(self, doc, judul: str) -> str:
        self.g += 1
        no = f"{self.bab}.{self.g}"
        p = doc.add_paragraph(style="_gambar")
        tulis_runs(p, f"Gambar {no} {judul}", size=Pt(11))
        self.daftar_g.append((no, judul))
        return no


def gambar(doc, nomor: Nomor, rel: str, judul: str, lebar_cm: float = 13.5) -> str:
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(0)
    p.add_run().add_picture(str(HB / rel), width=Cm(lebar_cm))
    return nomor.gambar(doc, judul)


def halaman_baru(doc):
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def main():
    from y26_bab4_isi import tulis_bab4
    from y26_bab45_lanjutan import tulis_bab4_lanjutan, tulis_bab5
    from y26_lampiran import tulis_lampiran

    doc = siapkan_dokumen()
    N4, data = tulis_bab4(doc)
    tulis_bab4_lanjutan(doc, N4, data)
    tulis_bab5(doc, data)
    tulis_lampiran(doc, N4)
    doc.save(OUT)
    print(f"OK -> {OUT}")
    print(f"   {len(N4.daftar_t)} tabel, {len(N4.daftar_g)} gambar pada BAB IV")
    print(f"   paragraf: {len(doc.paragraphs)}, tabel total: {len(doc.tables)}")


if __name__ == "__main__":
    main()
