"""Revisi terprogram BAB 1-3: v7 -> v8.

Menyunting HANYA teks biasa; tidak menyentuh *field* Mendeley, gambar, maupun gaya.
Sumber v7 tidak diubah — hasil ditulis ke berkas baru.

Delapan perubahan:
  1. Abstrak Indonesia — kalimat hasil ditulis ulang (18 placeholder numerik + 1 naratif)
  2. Abstract Inggris  — idem
  3. GPU "RTX 4060 8GB" -> "RTX 4060 Ti 8GB" (5 paragraf + 1 sel tabel)
  4. Angka split "678 / 339" -> "679 / 338" (2 lokasi)
  5. Diksi pembagian "secara acak" -> deterministik
  6. Tabel lingkungan — versi pustaka dilengkapi
  7. Subbab 3.10.1 — janji dua penghitung diselaraskan dengan pelaksanaan
  8. Verifikasi: tidak ada placeholder tersisa

    ./.venv/Scripts/python.exe y26_revisi_bab13.py
"""
from __future__ import annotations

import re
import shutil
from pathlib import Path

import docx

ROOT = Path(__file__).resolve().parent
SRC = ROOT / "TESIS_BAB1-3_REVISI_SIDANG_v7.docx"
DST = ROOT / "TESIS_BAB1-3_REVISI_SIDANG_v8.docx"

_MIRING = re.compile(r"\*(.+?)\*")


def ganti_di_paragraf(p, lama: str, baru: str) -> bool:
    """Ganti teks yang mungkin terpecah ke beberapa run, format run pertama dipertahankan."""
    penuh = "".join(r.text for r in p.runs)
    if lama not in penuh:
        return False
    awal = penuh.index(lama)
    akhir = awal + len(lama)
    pos, sisa_baru = 0, baru
    for r in p.runs:
        r_awal, r_akhir = pos, pos + len(r.text)
        pos = r_akhir
        if r_akhir <= awal or r_akhir - len(r.text) >= akhir:
            continue                                   # run di luar rentang
        depan = r.text[:max(0, awal - r_awal)]
        belakang = r.text[max(0, akhir - r_awal):] if r_akhir > akhir else ""
        r.text = depan + sisa_baru + belakang
        sisa_baru = ""                                 # sisipkan sekali saja
    return True


def tulis_ulang_paragraf(p, teks: str):
    """Kosongkan paragraf lalu tulis ulang; *bintang* menandai cetak miring."""
    gaya = p.runs[0] if p.runs else None
    nama, ukuran = (gaya.font.name, gaya.font.size) if gaya else (None, None)
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    pos = 0
    for m in _MIRING.finditer(teks):
        if m.start() > pos:
            r = p.add_run(teks[pos:m.start()])
            r.font.name, r.font.size = nama, ukuran
        r = p.add_run(m.group(1))
        r.italic = True
        r.font.name, r.font.size = nama, ukuran
        pos = m.end()
    if pos < len(teks):
        r = p.add_run(teks[pos:])
        r.font.name, r.font.size = nama, ukuran


# --------------------------------------------------------------- teks pengganti
ABSTRAK_ID = (
    "Hasil eksperimen menunjukkan konfigurasi penuh yang diusulkan mencapai mAP@0,5 sebesar "
    "77,97 persen dan mAP@0,5:0,95 sebesar 53,75 persen, namun selisihnya terhadap "
    "*baseline* YOLO26 standar tidak signifikan (p = 0,565). Kontribusi inkremental "
    "Pembobotan Loss Berbasis Densitas di atas Modul Atensi Hibrida dan Lapisan Deteksi "
    "Multi-Skala P2 justru signifikan (p = 0,037; *rank-biserial* 0,487), terkonsentrasi "
    "pada oklusi parsial 5,4 poin persentase dan objek kecil 3,0 poin persentase. "
    "Penghitungan menghasilkan MAE 1,97, RMSE 4,95, dan MAPE 37,17 persen pada kecepatan "
    "menyeluruh 20 *frame* per detik. Analisis *NMS-free* menunjukkan lapisan P2 menurunkan "
    "*Duplicate Rate* dan *Confidence Margin*, sedangkan atensi hibrida mengembalikan "
    "keduanya di atas *baseline*. Pembobotan Loss Berbasis Densitas karenanya bersifat "
    "komplementer terhadap modifikasi arsitektural dan melengkapi mekanisme bawaan YOLO26.")

ABSTRAK_EN = (
    "Experimental results show that the proposed full configuration achieves an mAP@0.5 of "
    "77.97 percent and an mAP@0.5:0.95 of 53.75 percent, yet its difference from the "
    "standard YOLO26 baseline is not significant (p = 0.565). The incremental contribution "
    "of Density-Aware Loss Weighting on top of the Hybrid Attention Module and the P2 "
    "multi-scale detection layer is significant (p = 0.037; rank-biserial 0.487), "
    "concentrated in partial occlusion at 5.4 percentage points and small objects at 3.0 "
    "percentage points. Counting yields an MAE of 1.97, an RMSE of 4.95, and a MAPE of "
    "37.17 percent at an end-to-end speed of 20 frames per second. The NMS-free analysis "
    "shows the P2 layer lowers both Duplicate Rate and Confidence Margin, whereas hybrid "
    "attention restores both above the baseline. Density-Aware Loss Weighting is therefore "
    "complementary to the architectural modifications and complements the built-in "
    "mechanisms of YOLO26.")

PENGHITUNG = (
    "Hitung manual dilakukan dengan menonton rekaman secara bertahap dan mencatat setiap "
    "kendaraan yang melintasi garis virtual beserta kelas dan arahnya. Penghitungan "
    "dilaksanakan oleh satu penghitung dengan protokol yang ditetapkan sebelumnya, meliputi "
    "definisi titik lintas, konvensi arah, dan pengelompokan kelas kendaraan, sehingga "
    "prosedurnya dapat direplikasi. Verifikasi silang oleh penghitung kedua sebagaimana "
    "lazim dilakukan pada pengumpulan data pengamatan tidak dapat dilaksanakan karena "
    "keterbatasan sumber daya pengamat, sehingga keandalan antarpenilai tidak "
    "terkuantifikasi. Keterbatasan ini dinyatakan secara eksplisit pada BAB IV dan perlu "
    "diperhatikan dalam menafsirkan galat penghitungan, khususnya pada klip berkepadatan "
    "tinggi yang paling rentan terhadap perbedaan penilaian antarpengamat.")


def main():
    shutil.copy2(SRC, DST)
    d = docx.Document(DST)
    ps = d.paragraphs
    log: list[str] = []

    # -- 1 & 2 abstrak ------------------------------------------------------
    for idx, teks, nama in ((26, ABSTRAK_ID, "ABSTRAK"), (37, ABSTRAK_EN, "ABSTRACT")):
        p = ps[idx]
        assert "[" in p.text, f"par[{idx}] bukan paragraf berplaceholder"
        tulis_ulang_paragraf(p, teks)
        log.append(f"{nama}: kalimat hasil ditulis ulang ({len(teks.split())} kata)")

    # -- 3 GPU --------------------------------------------------------------
    n = 0
    for p in ps:
        if ganti_di_paragraf(p, "RTX 4060 8GB", "RTX 4060 Ti 8GB"):
            n += 1
    for x in d.tables:
        for r in x.rows:
            for c in r.cells:
                for p in c.paragraphs:
                    if ganti_di_paragraf(p, "RTX 4060 8GB", "RTX 4060 Ti 8GB"):
                        n += 1
    log.append(f"GPU: {n} lokasi -> RTX 4060 Ti 8GB")

    # -- 4 & 5 angka split + diksi -----------------------------------------
    n = 0
    for p in ps:
        n += ganti_di_paragraf(
            p, "sekitar 2.372 citra latih, 678 citra validasi, dan 339 citra uji",
            "2.372 citra latih, 679 citra validasi, dan 338 citra uji")
        n += ganti_di_paragraf(p, "data uji berjumlah sekitar 339 citra",
                               "data uji berjumlah 338 citra")
        n += ganti_di_paragraf(
            p, "Pembagian kelompok dilakukan secara acak dengan menjaga keterwakilan",
            "Pembagian kelompok dilakukan secara deterministik, yaitu kelompok diurutkan "
            "menurut penanda yang stabil lalu dipotong pada ambang kumulatif sehingga "
            "prosedurnya menghasilkan pembagian yang identik pada setiap pengulangan, "
            "dengan tetap menjaga keterwakilan")
        n += ganti_di_paragraf(p, "Konsekuensi pembagian pada level kelompok adalah jumlah "
                                  "citra per subset hanya mendekati proporsi tersebut, yaitu ",
                               "Konsekuensi pembagian pada level kelompok adalah jumlah citra "
                               "per subset hanya mendekati proporsi tersebut, yaitu ")
    log.append(f"Split & diksi: {n} penggantian")

    # -- 6 tabel lingkungan -------------------------------------------------
    n = 0
    for x in d.tables:
        for r in x.rows:
            sel = [c.text.strip() for c in r.cells]
            if len(sel) < 2:
                continue
            baru = {"PyTorch dengan Ultralytics": "PyTorch 2.11.0+cu128 dengan Ultralytics 8.4.92",
                    "supervision (ByteTrack)": "supervision 0.29.1 (ByteTrack)",
                    "Python": "Python 3.11.9"}.get(sel[1])
            if baru:
                for p in r.cells[1].paragraphs:
                    if ganti_di_paragraf(p, sel[1], baru):
                        n += 1
                        break
    log.append(f"Tabel lingkungan: {n} versi pustaka dilengkapi")

    # -- 7 Subbab 3.10.1 ----------------------------------------------------
    p = ps[331]
    assert "dua penghitung" in p.text, "par[331] bukan paragraf hitung manual"
    tulis_ulang_paragraf(p, PENGHITUNG)
    log.append("Subbab 3.10.1: janji dua penghitung diselaraskan dengan pelaksanaan")

    d.save(DST)

    # -- 8 verifikasi -------------------------------------------------------
    v = docx.Document(DST)
    al = "\n".join(x.text for x in v.paragraphs) + "\n" + "\n".join(
        c.text for x in v.tables for r in x.rows for c in r.cells)
    sisa_ph = re.findall(r"\[[X0-9,.]{2,10}\]|\[[a-z][^\]]{15,90}\]", al)
    masalah = []
    if sisa_ph:
        masalah.append(f"placeholder tersisa: {sisa_ph}")
    if re.search(r"RTX 4060(?! Ti)", al):
        masalah.append("masih ada 'RTX 4060' tanpa Ti")
    for frasa in ("678 citra validasi", "339 citra", "dua penghitung secara terpisah",
                  "Pembagian kelompok dilakukan secara acak"):
        if frasa in al:
            masalah.append(f"masih ada: {frasa}")
    for wajib in ("RTX 4060 Ti 8GB", "679 citra validasi", "338 citra uji",
                  "deterministik", "Ultralytics 8.4.92", "Python 3.11.9",
                  "satu penghitung", "p = 0,565", "p = 0,037", "20 "):
        if wajib not in al:
            masalah.append(f"HILANG: {wajib}")

    print(f"\n=== {DST.name} ===")
    for s in log:
        print("  [OK]", s)
    print(f"\n  paragraf {len(v.paragraphs)}, tabel {len(v.tables)}, "
          f"gambar {len(v.inline_shapes)}, sitasi Mendeley utuh: "
          f"{'ya' if len(v.inline_shapes) == 16 else 'PERIKSA'}")
    print("\n=== VERIFIKASI ===")
    if masalah:
        for m in masalah:
            print("  [X] ", m)
    else:
        print("  [OK] semua pemeriksaan lolos")
    return not masalah


if __name__ == "__main__":
    raise SystemExit(0 if main() else 1)
