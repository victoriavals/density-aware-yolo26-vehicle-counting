#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Revisi naskah jurnal menurut KEPUTUSAN PEMBIMBING TEKNIS (II), 13 Agu 2026.

Bekerja di atas `JUTIF_Paper_..._REVISI_PROVENANS.docx` (hasil keputusan pertama).

  Bagian 1  Pembimbing MEMBATALKAN izin menarasikan tren p sebagai dukungan tafsir.
            Paragraf narasi monoton diganti rumusan yang beliau setujui: ketidaksepakatan
            antara uji peringkat-bertanda dan selang bootstrap **dilaporkan sebagai temuan**.
  K2        Bootstrap 1.000 -> **10.000 resample x 3 seed** (K-12 sudah tuntas 13,7 jam):
            Tabel 4, METHOD, dan seluruh penyebutan angka lama.
  D-G       50 citra Oculus New York dinyatakan BESERTA dampaknya pada pemilihan
            hiperparameter — validasi adalah subset yang dipakai grid search alfa/sigma
            DAN early stopping.
  D-H       Tanda air situs penjual dipisahkan kategorinya.
  D-I       381 citra potongan satu-kendaraan dipindah ke RESULT sebagai PENJELASAN
            sebaran sel, bukan sekadar keterbatasan.
  Temuan baru (audit K1 split uji): 14 thumbnail YouTube di split test, satu di antaranya
            memuat logo permainan berhak cipta; nol render dan nol citra bukan-lalu-lintas.

Naskah keputusan-pertama TIDAK diubah — keluaran berkas baru.
"""
from __future__ import annotations

import re
import shutil
import sys
from pathlib import Path

import docx

ROOT = Path(__file__).resolve().parent
SRC = ROOT / "JUTIF_Paper_DA-YOLO26_Firdaus_REVISI_PROVENANS.docx"
DST = ROOT / "JUTIF_Paper_DA-YOLO26_Firdaus_REVISI2.docx"

# ── Bagian 1: rumusan yang disetujui pembimbing, menggantikan narasi lama ──────
DISKUSI_BARU = (
    "On the full-configuration against baseline comparison the two analyses move in "
    "opposite directions as the test subset is cleaned. The signed-rank p value falls from "
    "0.565 to 0.303 and then to 0.079, while the bootstrap interval widens and begins to "
    "include zero, from [+0.0005, +0.0208] on the full subset to [−0.0005, +0.0305] on the "
    "CCTV-only subset. The divergence is explained by subset size rather than by effect: "
    "the CCTV-only subset contains 160 images, so the image-level resampling on which the "
    "interval is computed becomes markedly noisier while the per-cell paired ranks do not. "
    "Reporting only the p trend would therefore overstate the evidence. The consistent "
    "reading across both analyses and all three subsets remains that the full configuration "
    "is not reliably better than the baseline, and that the contribution of the weighting on "
    "top of the architectural instruments is."
)

# ── D-G: dampak komposisi validasi pada pemilihan hiperparameter ──────────────
VALIDASI = (
    "A further limitation concerns the validation subset rather than the test subset, and it "
    "bears on hyperparameter selection. Fifty of the 679 validation images, or 7.4 per cent, "
    "are stock footage of an indoor transit concourse that contains pedestrians and no "
    "vehicles at all; the same subset also holds 34 watermarked images and two rendered "
    "frames from driving-simulator software. The validation subset is what the grid search "
    "over the weighting strength and the density bandwidth was scored on, and it is also "
    "what early stopping monitored, so the single frozen setting used by all eight variants "
    "was selected on data that is not fully representative of the deployment condition. "
    "Repeating the search would require retraining and was not undertaken; the limitation is "
    "stated instead, and it offers a concrete explanation for an observation already reported "
    "here, namely that a single operating point tuned on one variant is not optimal for the "
    "others."
)

# ── D-I: penjelasan sebaran sel, ditempatkan di RESULT ───────────────────────
PENJELASAN_SEL = (
    "The composition of the dataset also explains a feature of the stratified cells that "
    "would otherwise look anomalous. Three families of catalogue imagery, comprising 161 "
    "weighbridge captures, 141 toll-classification captures and 79 roadside captures of "
    "tipper trucks, amount to 381 single-vehicle crops in which one large vehicle is "
    "photographed alone, unoccluded and at close range. Such images can only populate the "
    "large-size, unoccluded and sparse strata. In the test subset, 82 per cent of the "
    "ground-truth objects in the large-size big-vehicle cell come from catalogue stills "
    "rather than from surveillance frames. This is why the big-vehicle class contributes "
    "2,645 instances overall yet leaves only 17 objects in the small-size cell and 27 in the "
    "partially occluded cell, both of which fall below the minimum-cell rule and are excluded "
    "from testing. The imbalance is therefore a property of how the class was assembled, not "
    "a property of the detector."
)

# ── temuan audit K1 pada split uji ───────────────────────────────────────────
AUDIT_UJI = (
    "An exhaustive review of the 145 catalogue images in the test subset was carried out "
    "before submission, by perceptual-hash clustering followed by visual inspection of every "
    "cluster representative at native resolution. It found no rendered frames and no images "
    "without traffic content, so the reported metrics required no further recomputation. It "
    "did identify fourteen images whose file names are the default names of video-platform "
    "thumbnails and which carry channel branding, letterbox bars or title text, one of them "
    "additionally bearing the trademark of a commercial video game; these are photographs of "
    "real vehicles, so they affect attribution rather than construct validity, and they are "
    "counted with the third-party material declared below."
)

GANTI = [
    # Bagian 1 — ganti seluruh paragraf narasi monoton
    ("One pattern in the robustness analysis deserves a descriptive note.", None, DISKUSI_BARU,
     "Bagian 1: narasi monoton -> pelaporan ketidaksepakatan dua analisis"),
]

# K2 — penggantian angka bootstrap (teks lama -> teks baru), dicari di seluruh paragraf
GANTI_GLOBAL = [
    ("1,000 bootstrap resamples", "10,000 bootstrap resamples", "K2 METHOD"),
    ("1,000 resamples", "10,000 resamples", "K2 penyebutan umum"),
    ("bootstrap confidence intervals from 1,000", "bootstrap confidence intervals from 10,000",
     "K2 varian kalimat"),
    ("1000 resamples", "10,000 resamples", "K2 varian tanpa koma"),
]


def cari(doc, frasa):
    for p in doc.paragraphs:
        if frasa in p.text:
            return p
    return None


def tulis_ulang(p, teks):
    gaya = p.runs[0] if p.runs else None
    nama = gaya.font.name if gaya else None
    ukuran = gaya.font.size if gaya else None
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    r = p.add_run(teks)
    r.font.name, r.font.size = nama, ukuran


def ganti_di_paragraf(p, lama, baru):
    penuh = "".join(r.text for r in p.runs)
    if lama not in penuh:
        return False
    awal, akhir = penuh.index(lama), penuh.index(lama) + len(lama)
    pos, sisa = 0, baru
    for r in p.runs:
        ra, rb = pos, pos + len(r.text)
        pos = rb
        if rb <= awal or ra >= akhir:
            continue
        depan = r.text[:max(0, awal - ra)]
        belakang = r.text[max(0, akhir - ra):] if rb > akhir else ""
        r.text = depan + sisa + belakang
        sisa = ""
    return True


def sisip_setelah(doc, anchor, teks, style="BODY PARAGRAP"):
    baru = doc.add_paragraph(teks, style=style)
    anchor._p.addnext(baru._p)
    return baru


def main() -> int:
    if not SRC.exists():
        print(f"[gagal] {SRC.name} tidak ada")
        return 1
    shutil.copy2(SRC, DST)
    doc = docx.Document(str(DST))
    print("== Revisi jurnal — keputusan pembimbing teknis (II) ==\n")
    gagal = 0

    # 1. Bagian 1
    for kunci, _, baru, label in GANTI:
        p = cari(doc, kunci)
        if p is None:
            print(f"  [GAGAL] tak ditemukan: {kunci[:50]}")
            gagal += 1
            continue
        tulis_ulang(p, baru)
        print(f"  [ok] {label}")

    # 2. K2 — angka bootstrap di seluruh badan naskah + sel tabel
    n_k2 = 0
    for p in doc.paragraphs:
        for lama, baru, _ in GANTI_GLOBAL:
            if ganti_di_paragraf(p, lama, baru):
                n_k2 += 1
    for t in doc.tables:
        for row in t.rows:
            for c in row.cells:
                for p in c.paragraphs:
                    for lama, baru, _ in GANTI_GLOBAL:
                        if ganti_di_paragraf(p, lama, baru):
                            n_k2 += 1
    print(f"  [ok] K2: {n_k2} penyebutan bootstrap diperbarui ke 10.000")

    # 3. catatan kaki Tabel 4 — rentang lintas seed
    cap4 = cari(doc, "Table 4. Primary hypothesis tests")
    if cap4 is not None:
        sisip_setelah(doc, cap4,
                      "Bootstrap intervals are computed from 10,000 image-level resamples and "
                      "were repeated with three seeds. Across seeds the lower bound of the "
                      "V8-versus-V1 interval ranges from +0.0035 to +0.0216 percentage points, "
                      "so the interval excludes zero only marginally; the V8-versus-V5 interval "
                      "has a positive fraction of 1.000 on all three seeds, and the V4-versus-V1 "
                      "interval contains zero on all three.",
                      style="JUDUL TABEL")
        print("  [ok] K2: catatan rentang tiga seed di bawah keterangan Table 4")
    else:
        print("  [GAGAL] keterangan Table 4 tak ditemukan"); gagal += 1

    # 4. D-I — penjelasan sebaran sel di RESULT (sesudah subbab ketegaran)
    anchor = cari(doc, "Because part of the dataset is not surveillance imagery")
    if anchor is not None:
        a = sisip_setelah(doc, anchor, PENJELASAN_SEL)
        sisip_setelah(doc, a, AUDIT_UJI)
        print("  [ok] D-I penjelasan sebaran sel + temuan audit split uji (RESULT)")
    else:
        print("  [GAGAL] jangkar subbab ketegaran tak ditemukan"); gagal += 1

    # 5. D-G — dampak komposisi validasi pada penalaan (keterbatasan)
    anchor = cari(doc, "Two further limitations concern the provenance")
    if anchor is not None:
        sisip_setelah(doc, anchor, VALIDASI)
        print("  [ok] D-G komposisi validasi + dampak pada pemilihan hiperparameter")
    else:
        print("  [GAGAL] paragraf lisensi tak ditemukan"); gagal += 1

    # 6. D-H — pisahkan kategori tanda air situs penjual
    anchor = cari(doc, "Two further limitations concern the provenance")
    if anchor is not None:
        ok = ganti_di_paragraf(
            anchor,
            "Separately, four images are rendered frames produced by driving-simulator "
            "software rather than photographs, two of them in the validation subset.",
            "A separate category, distinct in origin from both, consists of commercial "
            "listing photographs carrying the watermark of the selling site or of a video "
            "channel; an exhaustive review of the test subset found such marks common there, "
            "and fourteen test images are video-platform thumbnails bearing channel branding. "
            "Finally, four images are rendered frames produced by driving-simulator software "
            "rather than photographs, two of them in the validation subset.")
        print(f"  [{'ok' if ok else 'GAGAL'}] D-H kategori tanda air situs dipisahkan")
        gagal += 0 if ok else 1

    if gagal:
        print(f"\n[BERHENTI] {gagal} langkah gagal — berkas keluaran TIDAK sahih.")
        return 1

    # ── verifikasi ───────────────────────────────────────────────────────────
    teks = "\n".join(p.text for p in doc.paragraphs)
    for t in doc.tables:
        for row in t.rows:
            for c in row.cells:
                teks += "\n" + c.text
    print("\n--- verifikasi ---")
    sisa = len(re.findall(r"1,?000\s+(bootstrap\s+)?resample", teks, re.I))
    print(f"  penyebutan '1.000 resample' tersisa : {sisa}  {'ok' if sisa == 0 else 'PERIKSA'}")
    print(f"  '10,000' hadir                      : {'ok' if '10,000' in teks else 'TIDAK'}")
    for frasa, label in (("move in opposite directions", "Bagian 1 rumusan pembimbing"),
                         ("82 per cent", "D-I angka 82 %"),
                         ("7.4 per cent", "D-G porsi Oculus"),
                         ("no rendered frames and no images", "temuan audit split uji"),
                         ("video-platform thumbnails", "D-H thumbnail YouTube")):
        print(f"  {label:32s}: {'ok' if frasa in teks else 'TIDAK ADA'}")
    lama = "One pattern in the robustness analysis"
    print(f"  narasi lama sudah hilang            : {'ok' if lama not in teks else 'MASIH ADA'}")
    ab = cari(doc, "Automatic vehicle counting from surveillance cameras")
    print(f"  abstrak                             : {len(ab.text.split())} kata (batas 250)")

    doc.save(str(DST))
    print(f"\nTersimpan: {DST.name}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
