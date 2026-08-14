#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Revisi naskah jurnal menurut KEPUTUSAN PEMBIMBING TEKNIS (III), 14 Agu 2026.

Bekerja di atas `..._REVISI2.docx` -> `..._SIAP_KIRIM.docx`.

  §2  Pernyataan keterbatasan diganti rumusan pembimbing yang **tegar terhadap angka akhir**
      ("at least 315", "may prove higher") karena audit 18 lembar sisa dinyatakan TIDAK
      memblokir pengiriman.
  §4  Temuan bias per sel (K4) dinaikkan dari pembelaan menjadi **klaim positif** — satu
      paragraf penuh di DISCUSSIONS.
  §5a Abstrak dipangkas 249 -> ~235 kata agar ada ruang bagi koreksi editorial penyunting.
  §5b Proporsi bagian: TIDAK diseimbangkan — terukur RESULT 48,7 %, di atas ambang 45 %.
  §5c Komposisi dataset dijadikan **Table 2**; Table 2..8 lama bergeser menjadi 3..9,
      termasuk seluruh rujukan silang di badan naskah.
  §6.6 Figure 9 diverifikasi sudah tersisip (versi utuh).

Penomoran ulang dikerjakan MENURUN (8->9, 7->8, ...) agar tidak bertabrakan, dan
diverifikasi setelahnya.
"""
from __future__ import annotations

import re
import shutil
import sys
from pathlib import Path

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

ROOT = Path(__file__).resolve().parent
SRC = ROOT / "JUTIF_Paper_DA-YOLO26_Firdaus_REVISI2.docx"
DST = ROOT / "JUTIF_Paper_DA-YOLO26_Firdaus_SIAP_KIRIM.docx"

TARGET_ABSTRAK = (228, 240)

# ── §2 pernyataan keterbatasan tegar (rumusan pembimbing, verbatim) ──────────
LISENSI_BARU = (
    "A provenance audit of the dataset identified at least 315 images carrying third-party "
    "watermarks, distributed as 248 in the training subset, 34 in validation and 33 in test. "
    "The evaluation subset was audited exhaustively and image by image; the audit of the "
    "1,597 web-sourced images in the training and validation subsets is ongoing, so the total "
    "may prove higher. No watermarked image is reproduced in any figure of this article, and "
    "the published version of the dataset excludes them. Of the 315, 311 originate from paid "
    "stock footage and four are frames of a cycling video channel, listed separately because "
    "channel material is not covered by any stock licence. A separate category consists of "
    "commercial listing photographs carrying the watermark of the selling site or of a video "
    "channel; an exhaustive review of the test subset found such marks common there, and "
    "fourteen test images are video-platform thumbnails bearing channel branding. Finally, "
    "four images are rendered frames produced by driving-simulator software rather than "
    "photographs, two of them in the validation subset."
)

# ── §4 temuan K4 sebagai KLAIM POSITIF di DISCUSSIONS ───────────────────────
KLAIM_K4 = (
    "A quantitative check on the dataset composition turns what would otherwise be a caveat "
    "into a positive result about the stratified claims. Catalogue stills are not spread "
    "evenly over the evaluation cells; they are concentrated almost entirely in three "
    "big-vehicle cells, supplying 81.8 per cent of the ground-truth objects in the large-size "
    "cell, 66.7 per cent in the sparse-density cell and 57.8 per cent in the unoccluded cell. "
    "The six cells that carry the incremental contribution of the weighting are by contrast "
    "measurably free of them: 2.6 and 1.9 per cent for partially occluded two-wheelers and "
    "cars, 1.3 and 1.2 per cent for small cars and pedestrians, 0.6 per cent for partially "
    "occluded pedestrians, and zero per cent for small two-wheelers, an average of 1.3 per "
    "cent against a mean of 11 per cent across all 24 tested cells. The advantage reported "
    "for this method therefore sits in cells that are almost entirely surveillance imagery, "
    "while the cells that catalogue stills dominate are precisely those that fall below the "
    "minimum-cell rule or support no claim at all. An imperfect dataset composition is thus "
    "shown, by measurement rather than by argument, not to contaminate the claims that rest "
    "on it."
)

# ── §5c Table 2 baru: komposisi dataset menurut sumber ──────────────────────
TABEL2_JUDUL = "Table 2. Dataset composition by source"
TABEL2 = [
    ["Source category", "Train", "Validation", "Test", "Total", "Share"],
    ["Web and catalogue stills (not surveillance)", "1,124", "328", "145", "1,597", "47.1%"],
    ["Indonesian traffic cameras, screen-recorded", "950", "202", "27", "1,179", "34.8%"],
    ["Stock footage, watermarked (Mecca, car-free day)", "229", "0", "0", "229", "6.8%"],
    ["Traffic cameras, Yogyakarta city ATCS", "0", "65", "96", "161", "4.8%"],
    ["Traffic cameras, Demak Regency", "50", "0", "37", "87", "2.6%"],
    ["Stock footage, watermarked (night arterial)", "0", "34", "33", "67", "2.0%"],
    ["Stock footage, unwatermarked (transit concourse)", "0", "50", "0", "50", "1.5%"],
    ["Stock footage, watermarked (Seoul)", "15", "0", "0", "15", "0.4%"],
    ["Video-channel material (cycling channel)", "4", "0", "0", "4", "0.1%"],
    ["Total", "2,372", "679", "338", "3,389", "100%"],
]
TABEL2_CATATAN = (
    "Indonesian traffic-camera imagery accounts for 1,427 images or 42.1 per cent of the "
    "total; stock footage of all kinds accounts for 365 or 10.8 per cent. The categories are "
    "assigned by file-name provenance reinforced by perceptual-hash clustering and visual "
    "inspection of every cluster representative."
)

# ── §5a pemangkasan abstrak ─────────────────────────────────────────────────
PANGKAS_ABSTRAK = [
    ("Differences were tested per class and per stratum of size, occlusion and density "
     "using Wilcoxon signed-rank tests with Holm correction, rank-biserial effect sizes "
     "and bootstrap intervals.",
     "Differences were tested per class and per stratum with Wilcoxon signed-rank tests, "
     "effect sizes and bootstrap intervals."),
    ("Density-Aware Loss Weighting did not improve the baseline alone and was significantly "
     "harmful on vehicle-only cells, yet on top of the architectural instruments its "
     "contribution was significant across all four classes",
     "Alone it did not improve the baseline and was significantly harmful on vehicle-only "
     "cells, yet on top of the architectural instruments its contribution was significant"),
    ("at 20.47 frames per second, below the 30 frames per second source rate.",
     "at 20.47 frames per second, below the 30 Hz source rate."),
]


def cari(doc, frasa):
    for p in doc.paragraphs:
        if frasa in p.text:
            return p
    return None


def ganti_di_paragraf(p, lama, baru):
    penuh = "".join(r.text for r in p.runs)
    if lama not in penuh:
        return False
    a, b = penuh.index(lama), penuh.index(lama) + len(lama)
    pos, sisa = 0, baru
    for r in p.runs:
        ra, rb = pos, pos + len(r.text)
        pos = rb
        if rb <= a or ra >= b:
            continue
        r.text = r.text[:max(0, a - ra)] + sisa + (r.text[max(0, b - ra):] if rb > b else "")
        sisa = ""
    return True


def tulis_ulang(p, teks):
    g = p.runs[0] if p.runs else None
    nama, uk = (g.font.name, g.font.size) if g else (None, None)
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    r = p.add_run(teks)
    r.font.name, r.font.size = nama, uk


def sisip_setelah(doc, anchor, teks, style="BODY PARAGRAP"):
    b = doc.add_paragraph(teks, style=style)
    anchor._p.addnext(b._p)
    return b


def sisip_tabel(doc, anchor, data):
    t = doc.add_table(rows=len(data), cols=len(data[0]))
    t.style = "Table Grid"
    for i, baris in enumerate(data):
        for j, sel in enumerate(baris):
            c = t.cell(i, j)
            c.text = ""
            p = c.paragraphs[0]
            p.style = doc.styles["BODY PARAGRAP"]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if j else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(sel)
            r.font.name, r.font.size = "Times New Roman", Pt(9)
            if i == 0 or baris[0] == "Total":
                r.bold = True
    anchor._p.addnext(t._tbl)
    return t


def semua_paragraf(doc):
    """Paragraf badan + paragraf di dalam sel tabel."""
    yield from doc.paragraphs
    for t in doc.tables:
        for row in t.rows:
            for c in row.cells:
                yield from c.paragraphs


def main() -> int:
    if not SRC.exists():
        print(f"[gagal] {SRC.name} tidak ada")
        return 1
    shutil.copy2(SRC, DST)
    doc = docx.Document(str(DST))
    print("== Revisi jurnal — keputusan pembimbing teknis (III) ==\n")
    gagal = 0

    # ── §5c penomoran ulang tabel MENURUN, sebelum Table 2 baru disisipkan ──
    print("-- §5c penomoran ulang Table 2..8 -> 3..9 --")
    n_ubah = 0
    for lama in range(8, 1, -1):
        baru = lama + 1
        for p in semua_paragraf(doc):
            penuh = "".join(r.text for r in p.runs)
            if f"Table {lama}" not in penuh:
                continue
            # ganti seluruh kemunculan 'Table N' pada paragraf ini
            while ganti_di_paragraf(p, f"Table {lama}", f"Table {baru}"):
                n_ubah += 1
                if f"Table {lama}" not in "".join(r.text for r in p.runs):
                    break
    print(f"  {n_ubah} rujukan/keterangan digeser")

    # ── §5c sisipkan Table 2 baru sesudah keterangan Table 1 ────────────────
    anchor = cari(doc, "Table 1. Dataset composition after group-based splitting")
    if anchor is None:
        print("  [GAGAL] keterangan Table 1 tak ditemukan"); return 1
    # lewati tabel Table 1 itu sendiri: sisipkan sesudah paragraf komposisi D-D
    komp = cari(doc, "The composition by source is less uniform")
    jangkar = komp if komp is not None else anchor
    cap = sisip_setelah(doc, jangkar, TABEL2_JUDUL, "JUDUL TABEL")
    sisip_tabel(doc, cap, TABEL2)
    # catatan di bawah tabel
    tbl_el = cap._p.getnext()
    catatan = doc.add_paragraph(TABEL2_CATATAN, style="BODY PARAGRAP")
    tbl_el.addnext(catatan._p)
    print(f"  [ok] Table 2 baru disisipkan ({len(TABEL2)-1} baris data)")

    # rujukan silang untuk Table 2 baru
    if not ganti_di_paragraf(komp, "is stated here in full because it bears directly",
                             "is stated here in full, and in Table 2, because it bears directly"):
        print("  [catatan] rujukan Table 2 tidak disisipkan di paragraf komposisi")

    # ── §2 pernyataan keterbatasan tegar ────────────────────────────────────
    p = cari(doc, "Two further limitations concern the provenance")
    if p is None:
        print("  [GAGAL] paragraf lisensi tak ditemukan"); gagal += 1
    else:
        tulis_ulang(p, LISENSI_BARU)
        print("  [ok] §2 pernyataan keterbatasan tegar ('at least' / 'may prove higher')")

    # ── §4 klaim K4 di DISCUSSIONS ──────────────────────────────────────────
    p = cari(doc, "On the full-configuration against baseline comparison")
    if p is None:
        print("  [GAGAL] jangkar DISCUSSIONS tak ditemukan"); gagal += 1
    else:
        sisip_setelah(doc, p, KLAIM_K4)
        print("  [ok] §4 temuan K4 dinaikkan menjadi paragraf klaim di DISCUSSIONS")

    # ── §5a pangkas abstrak ─────────────────────────────────────────────────
    ab = cari(doc, "Automatic vehicle counting from surveillance cameras")
    sebelum = len(ab.text.split())
    for lama, baru in PANGKAS_ABSTRAK:
        if not ganti_di_paragraf(ab, lama, baru):
            print(f"  [catatan] potongan abstrak tak cocok: {lama[:44]}…")
    sesudah = len(ab.text.split())
    lo, hi = TARGET_ABSTRAK
    print(f"  [ok] §5a abstrak {sebelum} -> {sesudah} kata "
          f"(sasaran {lo}-{hi}) {'OK' if lo <= sesudah <= hi else 'PERIKSA'}")

    if gagal:
        print(f"\n[BERHENTI] {gagal} langkah gagal.")
        return 1
    doc.save(str(DST))

    # ── verifikasi ──────────────────────────────────────────────────────────
    doc = docx.Document(str(DST))
    teks = "\n".join(p.text for p in semua_paragraf(doc))
    print("\n--- verifikasi ---")
    cap_t = sorted({int(m.group(1)) for m in re.finditer(r"^Table (\d+)\.", teks, re.M)})
    cap_f = sorted({int(m.group(1)) for m in re.finditer(r"^Figure (\d+)\.", teks, re.M)})
    print(f"  keterangan Table  : {cap_t}  {'OK' if cap_t == list(range(1, 10)) else 'PERIKSA'}")
    print(f"  keterangan Figure : {cap_f}  {'OK' if cap_f == list(range(1, 10)) else 'PERIKSA'}")
    print(f"  objek tabel       : {len(doc.tables)} (8 tabel isi + Table 9 + daftar periksa)")
    for n in range(1, 10):
        if f"Table {n}" not in teks:
            print(f"  ! Table {n} tidak dirujuk")
    for frasa, label in (("at least 315", "§2 'at least'"),
                         ("may prove higher", "§2 'may prove higher'"),
                         ("by measurement rather than by argument", "§4 klaim K4"),
                         ("81.8 per cent", "§4 angka konsentrasi"),
                         ("Dataset composition by source", "§5c Table 2")):
        print(f"  {label:24s}: {'ok' if frasa in teks else 'TIDAK ADA'}")
    ab = cari(doc, "Automatic vehicle counting")
    print(f"  abstrak           : {len(ab.text.split())} kata")
    print(f"\nTersimpan: {DST.name}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
