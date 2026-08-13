#!/usr/bin/env python
"""Revisi naskah JURNAL menurut keputusan pembimbing teknis 13 Agu 2026.

Menjalankan butir 3 dan 4 urutan pembimbing pada `JUTIF_Paper_..._FINAL.docx`:

  D-B   lokasi: Yogyakarta, Demak, Banjarmasin + pernyataan rekaman stok asing;
        angka kendaraan khusus Jakarta di INTRODUCTION dihapus (tidak diganti angka
        wilayah lain karena tidak ada sumber terverifikasi — CLAUDE.md §12.3).
  D-C   sumber data: "self-collected"/"data primer"/"kamera dipasang peneliti" dihapus;
        diganti deskripsi benar + atribusi lembaga (Method 2.1, keterangan Figure 9,
        dan ACKNOWLEDGEMENT).
  D-D   paragraf komposisi di Method 2.1, WAJIB memuat angka 82 % sel
        `size/large/big-vehicle` yang berasal dari citra katalog.
  Bag.3 subbab ketegaran di RESULT (Table 8, tiga subset berdampingan) + narasi pola
        p monoton di DISCUSSIONS **secara deskriptif tanpa klaim signifikansi**, beserta
        peringatan bahwa selang bootstrap pada subset terkecil melebar memuat nol.
  Bag.5 pernyataan lisensi: 4 citra kanal "NL Cycling" **dipisahkan** dari stok
        berlisensi (kedudukan hukum berbeda); kalimat hasil uji pHash; render simulator.
  D-A   Figure 9 versi PENUH (bukan ter-*crop*) + keterangan lima unsur.

Naskah asli TIDAK diubah — keluaran berkas baru. Naskah **tesis TIDAK disentuh**
(butir 5 menunggu arahan prosedur Dr. Sandfreni).

    python revisi_jurnal_provenans.py
    python revisi_jurnal_provenans.py --periksa-saja     # hanya laporkan, tanpa menulis
"""
from __future__ import annotations

import argparse
import re
import shutil
import sys
from pathlib import Path

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

ROOT = Path(__file__).resolve().parent
SRC = ROOT / "JUTIF_Paper_DA-YOLO26_Firdaus_FINAL.docx"
DST = ROOT / "JUTIF_Paper_DA-YOLO26_Firdaus_REVISI_PROVENANS.docx"
GAMBAR9 = ROOT / "hasil_bab4_5" / "15_gambar_jurnal" / "gambar_banding_en.png"

BATAS_ABSTRAK = 250          # daftar periksa JUTIF: 150-250 kata

# ───────────────────────────────────────────────────── penggantian dalam paragraf
# (kunci_paragraf, teks_lama, teks_baru, label)
GANTI: list[tuple[str, str, str, str]] = [
    # ── D-B + D-C pada ABSTRACT ────────────────────────────────────────────────
    ("Automatic vehicle counting from surveillance cameras",
     "A full factorial ablation of eight variants was trained on 3,389 self-collected "
     "CCTV images from Jakarta covering four classes, using group-based splitting to "
     "prevent leakage.",
     "A full factorial ablation of eight variants was trained on 3,389 images of four "
     "classes drawn from public traffic-camera feeds in Yogyakarta, Demak and "
     "Banjarmasin together with vehicle catalogue stills and licensed stock footage, "
     "using group-based splitting to prevent leakage.",
     "D-B/D-C abstrak: lokasi + hapus self-collected"),
    # kompresi agar abstrak tetap <= 250 kata
    ("Automatic vehicle counting from surveillance cameras",
     "Detection quality was evaluated per class and per stratum of size, occlusion and "
     "density, and differences were tested with Wilcoxon signed-rank tests, Holm "
     "correction, rank-biserial effect sizes and bootstrap confidence intervals.",
     "Differences were tested per class and per stratum of size, occlusion and density "
     "using Wilcoxon signed-rank tests with Holm correction, rank-biserial effect sizes "
     "and bootstrap intervals.",
     "kompresi abstrak (batas 250 kata)"),
    ("Automatic vehicle counting from surveillance cameras",
     "In this setting, the instrumentation shows that the NMS-free paradigm relocates "
     "rather than removes the duplication problem.",
     "The instrumentation shows that the NMS-free paradigm relocates rather than "
     "removes duplication.",
     "kompresi abstrak (batas 250 kata)"),

    # ── D-B pada INTRODUCTION: hapus angka khusus Jakarta ─────────────────────
    ("Traffic congestion in large Indonesian cities",
     " The same registry lists about 24.5 million vehicles in the jurisdiction of the "
     "Jakarta Metropolitan Police, close to 14.6% of the national fleet, of which about "
     "19.5 million are motorcycles [1].",
     "",
     "D-B introduction: hapus angka kendaraan khusus Jakarta"),

    # ── D-B + D-C pada Method 2.1 ─────────────────────────────────────────────
    ("The dataset consists of 3,389 images",
     "The dataset consists of 3,389 images recorded from urban traffic surveillance "
     "cameras in Jakarta and annotated into four classes, namely two-wheeler, car, big "
     "vehicle covering buses and trucks, and pedestrian as a non-vehicle context class "
     "[38].",
     "The dataset consists of 3,389 images annotated into four classes, namely "
     "two-wheeler, car, big vehicle covering buses and trucks, and pedestrian as a "
     "non-vehicle context class [38]. The surveillance imagery was captured by screen "
     "recording from publicly accessible traffic-camera feeds operated by three "
     "Indonesian local authorities, namely the area traffic control system of the "
     "Yogyakarta City Government, whose connectivity is credited on screen to corporate "
     "social responsibility partners, the Demak Regency Transport Agency, and the "
     "Banjarmasin City Transport Agency; the operating agencies are acknowledged "
     "accordingly. A smaller group of frames comes from consumer-grade fixed cameras. "
     "No camera was installed for this study, and none of the material is primary data "
     "collected by the authors.",
     "D-B/D-C method 2.1: lokasi + deskripsi sumber + atribusi lembaga"),

    # ── Bagian 5 butir 1: kalimat hasil uji pHash ─────────────────────────────
    ("The dataset consists of 3,389 images",
     "The procedure uses no random number generator, so it reproduces exactly.",
     "The procedure uses no random number generator, so it reproduces exactly. Because "
     "byte-level comparison cannot detect the same picture republished at a different "
     "size, the final split was additionally screened with a 64-bit perceptual hash "
     "over every pair of subsets. Three pairs of visually identical images with "
     "differing file hashes were found, one of them between the training and the test "
     "subset; that test image is excluded from the results reported here, while the "
     "training subset was left untouched because the weights had already been fitted "
     "with it.",
     "Bag.5: kalimat hasil uji pHash"),

    # ── D-B pada keterbatasan: "one city" -> tiga kota ────────────────────────
    ("Several limitations bound these conclusions",
     "Finally, the data covers one city, one camera viewpoint style, and a test subset "
     "dominated by night scenes,",
     "Finally, the surveillance data covers three cities and a small number of viewpoint "
     "styles, the test subset is dominated by night scenes,",
     "D-B keterbatasan: hapus klaim satu kota"),

    # ── D-C pada ACKNOWLEDGEMENT: atribusi lembaga ────────────────────────────
    ("The authors thank the Master of Computer Science",
     "for the academic support provided during this research.",
     "for the academic support provided during this research. The authors also "
     "acknowledge the area traffic control system of the Yogyakarta City Government, "
     "the Demak Regency Transport Agency and the Banjarmasin City Transport Agency, "
     "whose publicly accessible traffic-camera feeds provided the surveillance imagery "
     "analysed in this study.",
     "D-C acknowledgement: atribusi lembaga penyedia umpan"),
]

# ───────────────────────────────────────────────────────── paragraf/blok baru
KOMPOSISI = (
    "The composition by source is less uniform than the composition by subset, and it is "
    "stated here in full because it bears directly on how the stratified results should "
    "be read. Of the 3,389 images, 1,427 or 42.1% are Indonesian traffic-camera imagery, "
    "1,597 or 47.1% are vehicle catalogue, dealer-listing and news stills rather than "
    "surveillance frames, and 365 or 10.8% are stock footage, part of it recorded outside "
    "Indonesia. The catalogue stills were added to enlarge the big-vehicle class, which is "
    "the least frequent vehicle class in the surveillance material. The consequence for "
    "stratification is direct: a catalogue still shows a single large vehicle photographed "
    "from close range, so such images populate the large-size, unoccluded and sparse "
    "strata almost exclusively. In the test subset, 82% of the ground-truth objects in the "
    "large-size big-vehicle cell come from catalogue stills rather than from surveillance "
    "frames. That is also why the small-size big-vehicle cell holds only 17 objects even "
    "though the class contributes 2,645 instances overall, and it is the reason the "
    "big-vehicle cells behave differently from the rest. A robustness analysis that "
    "repeats the primary tests on surveillance-only and watermark-free test subsets is "
    "reported with the results."
)

KETEGARAN_JUDUL = "Robustness to Dataset Composition"
KETEGARAN_TEKS = (
    "Because part of the dataset is not surveillance imagery and part of it carries "
    "third-party watermarks, the three primary hypothesis tests were repeated on two "
    "reduced test subsets. The first removes the 33 watermarked images together with the "
    "one image identified by the perceptual-hash screen, leaving 304 images. The second "
    "keeps only the surveillance imagery, leaving 160 images. Neither reduction required "
    "retraining or repeated inference, because the stored per-image predictions of the "
    "one-to-one head were filtered by file name; as a control, the unreduced subset "
    "reproduces the published p values and bootstrap intervals exactly. The minimum-cell "
    "rule leaves the same 24 test units in all three subsets, so the three columns are "
    "directly comparable. Table 8 reports the outcome. The third hypothesis remains "
    "supported in all three subsets and neither of the other two changes direction. The "
    "two stratum-level differences reported for that hypothesis also persist: "
    "proxy-flagged partial occlusion moves from +5.37 to +6.57 and +6.10 percentage "
    "points, and small objects from +3.02 to +3.08 and +3.22 percentage points, across "
    "the full, watermark-free and surveillance-only subsets respectively."
)
TABEL8_JUDUL = ("Table 8. Robustness of the primary hypothesis tests to dataset "
                "composition")
TABEL8 = [
    ["Quantity", "Full test subset (338)", "Watermark-free (304)", "Surveillance only (160)"],
    ["V8 vs V1, unadjusted p", "0.565", "0.303", "0.079"],
    ["V8 vs V1, rank-biserial r", "+0.140", "+0.247", "+0.413"],
    ["V8 vs V1, bootstrap 95% CI (pp)", "+0.05 to +2.08", "+0.09 to +2.22", "−0.05 to +3.05"],
    ["V4 vs V1, unadjusted p", "0.208", "0.252", "0.546"],
    ["V4 vs V1, rank-biserial r", "−0.300", "−0.273", "−0.147"],
    ["V8 vs V5, unadjusted p", "0.037", "0.040", "0.023"],
    ["V8 vs V5, rank-biserial r", "+0.487", "+0.480", "+0.527"],
    ["V8 vs V5, bootstrap 95% CI (pp)", "+1.26 to +3.53", "+0.94 to +3.28", "+0.75 to +3.71"],
    ["Test units after minimum-cell rule", "24", "24", "24"],
]

DISKUSI_MONOTON = (
    "One pattern in the robustness analysis deserves a descriptive note. The evidence for "
    "the full configuration against the baseline strengthens monotonically as catalogue "
    "stills are removed, with the unadjusted p falling from 0.565 to 0.303 and then to "
    "0.079 and the effect size rising from 0.140 to 0.247 and then to 0.413. That is "
    "consistent with the interpretation that catalogue stills dilute the stratified "
    "evaluation, because they contribute large, unoccluded and uncrowded objects, which "
    "are precisely the conditions the method is not aimed at. The pattern is reported as a "
    "description and not as support for the first hypothesis. The comparison remains "
    "unsupported at the 5% level in every subset, and on the smallest subset the bootstrap "
    "interval widens to include zero, so once image-level uncertainty is taken into "
    "account the evidence is not in fact stronger. Both analyses agree that no reliable "
    "improvement over the baseline has been demonstrated."
)

LISENSI = (
    "Two further limitations concern the provenance of the material rather than the "
    "method. Of the 3,389 images, 315 carry a visible third-party watermark. Of these, "
    "311 originate from paid stock footage, whereas four are frames of a cycling video "
    "channel and are listed separately because channel material is not covered by any "
    "stock licence. Two hundred and forty-eight of the 315 fall in the training subset, "
    "so the fitted weights are bound to material of that kind and a complete removal "
    "would require retraining all eight variants; the 34 validation and 33 test images "
    "concerned are excluded in the robustness analysis reported above, none of the 315 "
    "appears in any figure of this article, and the published dataset record is being "
    "corrected. Separately, four images are rendered frames produced by driving-simulator "
    "software rather than photographs, two of them in the validation subset."
)

FIG9_RUJUKAN = (
    "Figure 9 illustrates the same trade-off on a single frame, with the ground truth "
    "shown alongside the predictions of the baseline and of the full configuration."
)
FIG9_KETERANGAN = (
    "Figure 9. Qualitative comparison on a single dense night frame "
    "(night-traffic-5_mp4-0028) from the ATCS feed of the Yogyakarta City Government at "
    "the Terban intersection, used for research purposes. Panels show (a) ground truth, "
    "(b) the baseline, and (c) the full configuration, at a confidence threshold of 0.25 "
    "and an intersection over union of 0.50. The full configuration recovers one "
    "additional object at the cost of one additional false positive, which is the same "
    "trade-off observed in the aggregate error analysis. A single frame is an "
    "illustration and not evidence; no quantitative claim rests on it."
)


# ───────────────────────────────────────────────────────────────── utilitas docx
def cari(doc, frasa: str):
    for p in doc.paragraphs:
        if frasa in p.text:
            return p
    return None


def ganti_di_paragraf(p, lama: str, baru: str) -> bool:
    """Ganti substring yang mungkin terpecah antar-run; format run pertama dipertahankan."""
    penuh = "".join(r.text for r in p.runs)
    if lama not in penuh:
        return False
    awal, akhir = penuh.index(lama), penuh.index(lama) + len(lama)
    pos, sisa = 0, baru
    for r in p.runs:
        r_awal, r_akhir = pos, pos + len(r.text)
        pos = r_akhir
        if r_akhir <= awal or r_awal >= akhir:
            continue
        depan = r.text[:max(0, awal - r_awal)]
        belakang = r.text[max(0, akhir - r_awal):] if r_akhir > akhir else ""
        r.text = depan + sisa + belakang
        sisa = ""
    return True


def sisip_paragraf(doc, anchor, teks: str, style: str, align=None):
    """Buat paragraf bergaya `style` lalu pindahkan tepat SESUDAH `anchor`."""
    baru = doc.add_paragraph(teks, style=style)
    if align is not None:
        baru.alignment = align
    anchor._p.addnext(baru._p)
    return baru


def sisip_tabel(doc, anchor, data: list[list[str]], style="Table Grid"):
    t = doc.add_table(rows=len(data), cols=len(data[0]))
    t.style = style
    for i, baris in enumerate(data):
        for j, sel in enumerate(baris):
            c = t.cell(i, j)
            c.text = ""
            p = c.paragraphs[0]
            p.style = doc.styles["BODY PARAGRAP"]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(sel)
            r.font.name = "Times New Roman"
            r.font.size = Pt(10)
            if i == 0:
                r.bold = True
    anchor._p.addnext(t._tbl)
    return t


def sisip_gambar(doc, anchor, path: Path, lebar_cm: float):
    p = doc.add_paragraph(style="JUDUL GAMBAR")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Cm(lebar_cm))
    anchor._p.addnext(p._p)
    return p


def kata(teks: str) -> int:
    return len(teks.split())


# ─────────────────────────────────────────────────────────────────────── main
def main() -> int:
    ap = argparse.ArgumentParser(description=__doc__,
                                 formatter_class=argparse.RawDescriptionHelpFormatter)
    ap.add_argument("--periksa-saja", action="store_true")
    ap.add_argument("--lebar-gambar-cm", type=float, default=16.0)
    args = ap.parse_args()

    if not SRC.exists():
        print(f"[gagal] {SRC.name} tidak ada")
        return 1
    if not GAMBAR9.exists():
        print(f"[gagal] {GAMBAR9} tidak ada — bangkitkan dulu y26_gambar_jurnal.py")
        return 1

    if args.periksa_saja:
        doc = docx.Document(str(SRC))
    else:
        shutil.copy2(SRC, DST)
        doc = docx.Document(str(DST))

    print("== Revisi naskah jurnal menurut keputusan pembimbing 13 Agu 2026 ==\n")

    # ── 1. penggantian teks ───────────────────────────────────────────────────
    gagal = 0
    for kunci, lama, baru, label in GANTI:
        p = cari(doc, kunci)
        if p is None:
            print(f"  [GAGAL] paragraf kunci tak ditemukan: {kunci[:44]!r}")
            gagal += 1
            continue
        if ganti_di_paragraf(p, lama, baru):
            print(f"  [ok] {label}")
        else:
            print(f"  [GAGAL] teks lama tak cocok — {label}")
            gagal += 1
    if gagal:
        print(f"\n[BERHENTI] {gagal} penggantian gagal; berkas keluaran TIDAK sahih.")
        return 1

    # ── 2. paragraf komposisi D-D, sesudah keterangan Table 1 ─────────────────
    # Ditempatkan sesudah tabel komposisi agar pembaca melihat angka per subset dulu.
    anchor = cari(doc, "Table 1. Dataset composition")
    sisip_paragraf(doc, anchor, KOMPOSISI, "BODY PARAGRAP")
    print("  [ok] D-D paragraf komposisi (memuat angka 82 %)")

    # ── 3. subbab ketegaran + Table 8, sesudah Figure 3 (Stratified) ──────────
    anchor = cari(doc, "Figure 3. Stratified average precision")
    a1 = sisip_paragraf(doc, anchor, KETEGARAN_JUDUL, "SUB JUDUL")
    a2 = sisip_paragraf(doc, a1, KETEGARAN_TEKS, "BODY PARAGRAP")
    a3 = sisip_paragraf(doc, a2, TABEL8_JUDUL, "JUDUL TABEL")
    sisip_tabel(doc, a3, TABEL8)
    print("  [ok] Bag.3 subbab ketegaran + Table 8 (tiga subset)")

    # ── 4. narasi monoton di DISCUSSIONS (deskriptif + peringatan bootstrap) ──
    anchor = cari(doc, "The central result is that density-aware weighting")
    sisip_paragraf(doc, anchor, DISKUSI_MONOTON, "BODY PARAGRAP")
    print("  [ok] Bag.3 narasi pola p monoton (deskriptif, + peringatan bootstrap)")

    # ── 5. pernyataan lisensi sesudah paragraf keterbatasan ───────────────────
    anchor = cari(doc, "Several limitations bound these conclusions")
    sisip_paragraf(doc, anchor, LISENSI, "BODY PARAGRAP")
    print("  [ok] Bag.5 pernyataan lisensi (NL Cycling dipisahkan) + render simulator")

    # ── 6. Figure 9 sesudah keterangan Figure 8 (akhir Error Analysis) ───────
    anchor = cari(doc, "Figure 8. Missed objects by illumination")
    r1 = sisip_paragraf(doc, anchor, FIG9_RUJUKAN, "BODY PARAGRAP")
    g = sisip_gambar(doc, r1, GAMBAR9, args.lebar_gambar_cm)
    sisip_paragraf(doc, g, FIG9_KETERANGAN, "JUDUL GAMBAR", WD_ALIGN_PARAGRAPH.CENTER)
    print(f"  [ok] D-A Figure 9 versi PENUH ({args.lebar_gambar_cm} cm) + keterangan lima unsur")

    # ── 7. verifikasi ─────────────────────────────────────────────────────────
    abstrak = cari(doc, "Automatic vehicle counting from surveillance cameras")
    n_kata = kata(abstrak.text)
    print(f"\n--- verifikasi ---")
    print(f"  abstrak: {n_kata} kata  (batas JUTIF 150-250)  "
          f"{'LOLOS' if 150 <= n_kata <= BATAS_ABSTRAK else 'GAGAL'}")
    teks_semua = "\n".join(p.text for p in doc.paragraphs)
    terlarang = {
        "self-collected": "klaim D-C",
        "data primer": "klaim D-C",
        "dipasang peneliti": "klaim D-C",
        "one city": "klaim D-B",
    }
    for frasa, ket in terlarang.items():
        n = len(re.findall(re.escape(frasa), teks_semua, re.I))
        print(f"  frasa terlarang {frasa!r:22s}: {n}  "
              f"{'ok' if n == 0 else 'MASIH ADA — ' + ket}")
    jak = [p.text[:70] for p in doc.paragraphs
           if re.search(r"(?i)\bjakarta\b", p.text)
           and not re.search(r"(?i)jakarta.cikampek", p.text)]
    print(f"  paragraf ber-'Jakarta' selain sitasi [34]: {len(jak)}"
          + ("".join(f"\n      ! {t}" for t in jak) if jak else "  ok"))
    print(f"  wajib ada '82%' (D-D): "
          f"{'ok' if '82%' in teks_semua else 'TIDAK ADA'}")
    print(f"  wajib ada kalimat 'not evidence' (D-A3): "
          f"{'ok' if 'illustration and not evidence' in teks_semua else 'TIDAK ADA'}")
    print(f"  tabel: {len(doc.tables)} (semula 8, harap 9)")
    n_gambar = len(re.findall(r"<w:drawing", doc.element.body.xml))
    print(f"  gambar: {n_gambar} (semula 8, harap 9)")

    if not args.periksa_saja:
        doc.save(str(DST))
        print(f"\nTersimpan: {DST.name}   (naskah asli TIDAK diubah)")
        print("Naskah TESIS tidak disentuh — butir 5 menunggu arahan Dr. Sandfreni.")
    return 0


if __name__ == "__main__":
    sys.exit(main())
